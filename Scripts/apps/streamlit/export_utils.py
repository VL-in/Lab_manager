"""Exportação de conversas (DOCX)."""

from __future__ import annotations

from io import BytesIO
from typing import Any

try:
    from docx import Document
except ImportError:  # pragma: no cover
    Document = None  # type: ignore[misc, assignment]


def conversation_to_docx_bytes(messages: list[dict[str, Any]], title: str = "Exportação") -> bytes:
    if Document is None:
        raise RuntimeError("O pacote python-docx não está instalado.")
    doc = Document()
    doc.add_heading("Assistente de laboratório — conversa exportada", level=0)
    doc.add_paragraph(title, style="Intense Quote")
    for m in messages:
        role = m.get("role", "")
        label = "Usuário" if role == "user" else "Assistente"
        doc.add_heading(label, level=2)
        doc.add_paragraph(str(m.get("content", "")))
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
